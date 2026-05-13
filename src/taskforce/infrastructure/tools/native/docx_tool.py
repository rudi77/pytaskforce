"""DOCX document manipulation tool.

Provides reading, writing, and editing capabilities for Microsoft Word
(.docx) files. ``python-docx`` ships as a core dependency.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from taskforce.core.interfaces.tools import ApprovalRiskLevel
from taskforce.infrastructure.tools.base_tool import BaseTool

logger = structlog.get_logger(__name__)


def _require_docx() -> Any:
    """Lazily import python-docx and raise a clear error if missing."""
    try:
        import docx  # noqa: F811

        return docx
    except ImportError:
        raise ImportError(
            "python-docx is part of the core install but the import failed. "
            "Run 'uv sync' to repair the venv."
        )


class DocxTool(BaseTool):
    """Read, create, and edit Microsoft Word (.docx) documents."""

    tool_name = "docx"
    tool_description = (
        "Manipulate Microsoft Word (.docx) files. "
        "Supports reading text, creating new documents, adding/replacing paragraphs, "
        "adding tables, and listing document structure."
    )
    tool_parameters_schema = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": ["read", "create", "add_paragraph", "replace_text", "add_table", "info"],
                "description": (
                    "Action to perform: "
                    "'read' extracts all text, "
                    "'create' creates a new document, "
                    "'add_paragraph' appends a paragraph, "
                    "'replace_text' finds and replaces text, "
                    "'add_table' adds a table, "
                    "'info' returns document metadata and structure."
                ),
            },
            "path": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "content": {
                "type": "string",
                "description": "Text content (for create/add_paragraph).",
            },
            "style": {
                "type": "string",
                "description": "Paragraph or heading style (e.g. 'Heading 1', 'Normal').",
            },
            "find": {
                "type": "string",
                "description": "Text to find (for replace_text).",
            },
            "replace": {
                "type": "string",
                "description": "Replacement text (for replace_text).",
            },
            "rows": {
                "type": "array",
                "items": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "description": "Table data as list of rows, each row a list of cell strings.",
            },
            "headers": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Optional header row for add_table.",
            },
        },
        "required": ["action", "path"],
    }

    tool_requires_approval = True
    tool_approval_risk_level = ApprovalRiskLevel.MEDIUM
    tool_supports_parallelism = True

    async def _execute(self, **kwargs: Any) -> dict[str, Any]:
        """Dispatch to the requested action handler."""
        action = kwargs["action"]
        handler = {
            "read": self._read,
            "create": self._create,
            "add_paragraph": self._add_paragraph,
            "replace_text": self._replace_text,
            "add_table": self._add_table,
            "info": self._info,
        }.get(action)

        if handler is None:
            return {"success": False, "error": f"Unknown action: {action}"}

        return await handler(**kwargs)

    async def _read(self, **kwargs: Any) -> dict[str, Any]:
        """Read all text from a .docx file."""
        docx = _require_docx()
        path = kwargs["path"]
        doc = docx.Document(path)

        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        return {
            "success": True,
            "text": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }

    async def _create(self, **kwargs: Any) -> dict[str, Any]:
        """Create a new .docx document with optional content."""
        docx = _require_docx()
        path = kwargs["path"]
        content = kwargs.get("content", "")
        style = kwargs.get("style")

        doc = docx.Document()
        if content:
            if style:
                doc.add_paragraph(content, style=style)
            else:
                doc.add_paragraph(content)

        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        logger.info("docx.created", path=path)
        return {"success": True, "path": path, "message": f"Document created at {path}"}

    async def _add_paragraph(self, **kwargs: Any) -> dict[str, Any]:
        """Append a paragraph to an existing .docx file."""
        docx = _require_docx()
        path = kwargs["path"]
        content = kwargs.get("content", "")
        style = kwargs.get("style")

        if not content:
            return {"success": False, "error": "Parameter 'content' is required for add_paragraph"}

        doc = docx.Document(path)
        if style:
            doc.add_paragraph(content, style=style)
        else:
            doc.add_paragraph(content)
        doc.save(path)

        logger.info("docx.paragraph_added", path=path)
        return {"success": True, "path": path, "message": "Paragraph added"}

    async def _replace_text(self, **kwargs: Any) -> dict[str, Any]:
        """Find and replace text across all paragraphs."""
        docx = _require_docx()
        path = kwargs["path"]
        find_text = kwargs.get("find", "")
        replace_text = kwargs.get("replace", "")

        if not find_text:
            return {"success": False, "error": "Parameter 'find' is required for replace_text"}

        doc = docx.Document(path)
        count = 0
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                for run in paragraph.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        count += 1
        doc.save(path)

        logger.info("docx.text_replaced", path=path, replacements=count)
        return {
            "success": True,
            "path": path,
            "replacements": count,
            "message": f"Replaced {count} occurrence(s)",
        }

    async def _add_table(self, **kwargs: Any) -> dict[str, Any]:
        """Add a table to an existing .docx file."""
        docx = _require_docx()
        path = kwargs["path"]
        rows = kwargs.get("rows", [])
        headers = kwargs.get("headers")

        if not rows and not headers:
            return {"success": False, "error": "Parameter 'rows' or 'headers' is required"}

        doc = docx.Document(path)

        all_rows = []
        if headers:
            all_rows.append(headers)
        all_rows.extend(rows)

        if not all_rows:
            return {"success": False, "error": "No data provided for table"}

        num_cols = len(all_rows[0])
        table = doc.add_table(rows=len(all_rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(all_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    table.rows[i].cells[j].text = str(cell_text)

        doc.save(path)
        logger.info("docx.table_added", path=path, rows=len(all_rows), cols=num_cols)
        return {
            "success": True,
            "path": path,
            "rows": len(all_rows),
            "cols": num_cols,
            "message": f"Table added ({len(all_rows)} rows x {num_cols} cols)",
        }

    async def _info(self, **kwargs: Any) -> dict[str, Any]:
        """Return document metadata and structure info."""
        docx = _require_docx()
        path = kwargs["path"]
        doc = docx.Document(path)

        styles_used = set()
        headings = []
        for p in doc.paragraphs:
            if p.style and p.style.name:
                styles_used.add(p.style.name)
                if p.style.name.startswith("Heading"):
                    headings.append({"level": p.style.name, "text": p.text})

        core_props = doc.core_properties
        return {
            "success": True,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "styles_used": sorted(styles_used),
            "headings": headings,
            "author": core_props.author or "",
            "title": core_props.title or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
